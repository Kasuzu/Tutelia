# ============================================================
# Wizard de Tutela con persistencia local (SQLite) y export a
# Word (.docx). Flujo: primero usuario diligencia; luego IA.
# SE CAMBIÓ:
# - Pruebas + Anexos -> una sola sección: "pruebas_y_anexos"
# - Mejora automática al guardar: hechos, pretensiones (con sugerencias) y pruebas_y_anexos
# - Cadena IA (un solo paso) desde hechos (+pretensiones): derechos_vulnerados → fundamentos_juridicos → fundamentos_de_derecho → ref
# - Notificaciones y Firmas se auto-componen desde "parties"
# - Pretensiones: prohibido contenido económico
# - Endpoint para ver texto final consolidado (/case/{id}/compose-final) y bundle estructurado (/case/{id}/compose-structured)
# ============================================================

import os
import json
import uuid
import sqlite3
import re
from datetime import datetime
from typing import Dict, Any, Optional, List, Tuple

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
import unicodedata
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ------------------------------------------------------------
# Config & Constantes
# ------------------------------------------------------------

DATA_DIR_DEFAULT = "./data"
EXPORT_DIR_DEFAULT = "./exports"

# Encabezado fijo
HEADER_FIXED = (
    "SEÑOR\n"
    "JUEZ DE LA REPÚBLICA (REPARTO)\n"
    "E. S. D.\n"
)

# Guías (LLM prompts) — versión actualizada
guides: Dict[str, str] = {
    "hechos": (
        "Redacta HECHOS en español claro, en orden cronológico, conciso y verificable. "
        "Devuelve una lista numerada (1), (2), (3)... sin encabezados ni preámbulos. "
        "Prohibido: frases sobre 'análisis', 'según lo anterior' o hablar de ti."
    ),
    "derechos_vulnerados": (
        "Redacta DERECHOS VULNERADOS conectados con los hechos. "
        "Primero nómbralos y luego explica brevemente el porqué. "
        "Usa viñetas o numeración; sin preámbulos ni metadiscurso. "
        "Cuando citemos normas o jurisprudencia, si usas texto literal ponlo entre comillas."
    ),
    "fundamentos_juridicos": (
        "Construye FUNDAMENTOS JURÍDICOS que conecten hechos, pruebas y reglas aplicables. "
        "Redáctalos en párrafos breves o lista numerada, sin metadiscurso ni citas extensas. "
        "Estructura sugerida: "
        "1) Procedencia: analiza subsidiariedad, inmediatez, legitimación por activa/pasiva y posible perjuicio irremediable, "
        "referenciando (H#) cuando sea necesario. "
        "2) Problema jurídico: formula una pregunta clara sobre la eventual vulneración del derecho (D#). "
        "3) Reglas jurisprudenciales y legales pertinentes (enunciadas de forma sintética). "
        "4) Caso concreto: subsume los hechos a las reglas y explica por qué se configura (o no) la vulneración. "
        "Evita conclusiones tajantes sin anclaje fáctico; sustenta siempre con (H#) y, cuando proceda, alude al soporte (P#)."
    ),
    "pretensiones": (
        "Organiza PRETENSIONES en lista numerada con órdenes claras y plazos razonables. "
        "Puedes sugerir complementarias. Sin metadiscurso."
        "No pongas pretensiones economicas o relacionadas al dinero"
    ),
    "pruebas_y_anexos": (
        "Limpia y ordena PRUEBAS Y ANEXOS en lista numerada. Sin preámbulos."
        "No inventes documentos , solo mejora la redaccion de los documentos presentados por el usuario"
    ),
    "fundamentos_de_derecho": (
        "Devuelve SOLO una lista numerada de normas y sentencias aplicables, sin explicar. "
        "Usa un formato de cita breve y estandarizado, por ejemplo: "
        "1) 'C.P., art. 86'; 2) 'D. 2591 de 1991, arts. 5, 6 y 42'; 3) 'Ley 1751 de 2015, art. 2'; "
        "4) 'CC, T-###/AAAA'; 5) 'CC, SU-###/AAAA'. "
        "Incluye únicamente disposiciones pertinentes al caso; no inventes números ni años. "
        "No agregues comentarios, glosas ni textos introductorios."
    ),
    "ref": (
    "Devuelve una sola línea de REFERENCIA precisa, sin preámbulos, con este formato: "
    "'Acción de tutela para la protección de los derechos fundamentales a [derechos] "
    "(con conexidad con [otros, si aplica]), interpuesta por [NOMBRES ACCIONANTES] "
    "contra [NOMBRES ACCIONADOS].' "
    "Usa exactamente los nombres que te paso en 'Accionantes' y 'Accionados'; "
    "NO uses corchetes ni marcadores. Termina en punto."
    ),

}

# Paso -> metadata (si requiere LLM y orden para pipeline manual).
# Nota: seguimos mejorando automáticamente al guardar: hechos, pretensiones, pruebas_y_anexos.
# La cadena jurídica se genera por-pantalla con endpoints ensure/*.
SECTIONS_CONFIG: Dict[str, Dict[str, Any]] = {
    "encabezado": {"needs_llm": False, "send_order": None},
    "accionantes": {"needs_llm": False, "send_order": None},
    "accionados": {"needs_llm": False, "send_order": None},
    "intro": {"needs_llm": False, "send_order": None},
    "hechos": {"needs_llm": True, "send_order": 1},
    "pruebas_y_anexos": {"needs_llm": False, "send_order": None},  # IA al guardar
    "pretensiones": {"needs_llm": False, "send_order": None},       # IA al guardar
    "derechos_vulnerados": {"needs_llm": True, "send_order": 2},
    "fundamentos_juridicos": {"needs_llm": True, "send_order": 3},
    "fundamentos_de_derecho": {"needs_llm": True, "send_order": 4},
    "ref": {"needs_llm": True, "send_order": 5},
    "cumplimiento_art_37": {"needs_llm": False, "send_order": None},
    "notificaciones": {"needs_llm": False, "send_order": None},
    "firmas": {"needs_llm": False, "send_order": None},
}

RIGHTS_LEXICON = {
    # Salud y conexos
    "salud": [
        "salud","derecho a la salud","ips","eps","ese","hospital","clinica","urgencias","triage",
        "cita","oportunidad en la atencion","remision","traslado","autorizacion","negacion del servicio",
        "medicamento","medicamentos","tratamiento","terapia","cirugia","procedimiento","examen",
        "orden medica","formula medica","historia clinica","mipres","pbs","pos","no pbs","no pos",
        "cups","cie10","rehabilitacion","consulta externa","glosa","barrera administrativa"
    ],
    # Vida y vida digna
    "vida digna": [
        "vida digna","existencia digna","indigno","condiciones inhumanas","tratos crueles",
        "riesgo vital","amenaza a la vida","supervivencia","dignidad humana"
    ],
    # Mínimo vital y subsistencia
    "minimo vital": [
        "minimo vital","subsistencia","alimento","alimentos","sustento","canasta basica",
        "ayuda humanitaria","auxilio","privacion de ingresos","no pago de salarios","pago atrasado"
    ],
    # Debido proceso (adm., judicial y disciplinario)
    "debido proceso": [
        "debido proceso","defensa","contradiccion","imparcialidad","legalidad",
        "procedimiento","actuacion administrativa","proceso disciplinario","sumario",
        "notificacion","notificado","traslado","termino","plazo","vencimiento de termino",
        "recursos","recurso de reposicion","apelacion","queja","nulidad","pruebas","audiencia",
        "motivacion","resolucion motivada","silencio administrativo"
    ],
    # Educación (acceso, permanencia, apoyos)
    "educacion": [
        "educacion","colegio","institucion educativa","universidad","matricula","matrícula",
        "cupos","traslado escolar","certificados","acta de grado","diploma","pae","transporte escolar",
        "ajustes razonables","inclusion educativa","docente de apoyo","material pedagogico"
    ],
    # Seguridad social (pensión, riesgos, licencias)
    "seguridad social": [
        "seguridad social","pension","pensión","colpensiones","rais","afp","arl","caja de compensacion",
        "cotizacion","cotizaciones","historia laboral","semanas cotizadas","licencia de maternidad",
        "incapacidad","subsidio familiar","ibc"
    ],
    # Petición y acceso a información (respuesta oportuna)
    "peticion": [
        "derecho de peticion","peticion","pqrs","pqr","pqrds","radicado","respuesta",
        "no respondio","no contesto","termino de 15 dias","quince dias","informacion solicitada",
        "entrega de informacion","trasparencia","ley 1755"
    ],
    # Igualdad y no discriminación
    "igualdad": [
        "igualdad","no discriminacion","discriminacion","trato desigual","enfoque diferencial",
        "discapacidad","enfoque de genero","orientacion sexual","identidad de genero","poblacion afro",
        "pueblos indigenas","adulto mayor","migrante"
    ],
    # Hábeas data (datos personales y financieros)
    "habeas data": [
        "habeas data","datos personales","proteccion de datos","actualizacion de datos",
        "rectificacion de datos","supresion de datos","autorizacion de datos",
        "historia crediticia","centrales de riesgo","datacredito","cifin","reporte negativo"
    ],
    # Libertad de expresión e información
    "libertad de expresion": [
        "libertad de expresion","censura","retiro de contenido","bloqueo de cuenta",
        "opinion","informacion veraz","rectificacion"
    ],
    # Trabajo y estabilidad reforzada
    "trabajo": [
        "trabajo","empleo","contrato laboral","despido","salario","prestaciones",
        "acoso laboral","estabilidad laboral reforzada","fuero de salud","fuero sindical",
        "reintegro","nomina","pago de horas"
    ],
    # Vivienda digna y servicios públicos domiciliarios
    "vivienda digna": [
        "vivienda digna","desalojo","reubicacion","techo","hacinamiento",
        "servicios publicos","energia","agua","gas","corte del servicio","facturacion",
        "suspension del servicio"
    ],
    # Agua potable y saneamiento básico (conexidad vida/vida digna)
    "agua potable": [
        "agua potable","acueducto","alcantarillado","saneamiento basico",
        "corte de agua","suspension de agua","carrotanque","potabilizacion"
    ],
    # Ambiente sano
    "ambiente sano": [
        "ambiente sano","contaminacion","ruido","emisiones","desechos","basuras",
        "licencia ambiental","impacto ambiental"
    ],
    # Libertad personal / Hábeas corpus
    "libertad personal": [
        "libertad personal","habeas corpus","retencion","captura","privacion de la libertad",
        "traslado carcelario","inpec","upj","demora en audiencia"
    ],
    # Familia, niñez y adolescencia
    "familia y niniez": [
        "familia","unidad familiar","custodia","visitas","icbf","comisaria de familia",
        "interes superior del menor","nna","proteccion integral","restablecimiento de derechos"
    ],
    # Personas con discapacidad (ajustes razonables)
    "personas con discapacidad": [
        "discapacidad","certificado de discapacidad","ajustes razonables","rehabilitacion",
        "ayudas tecnicas","silla de ruedas","accesibilidad","lengua de senas","lector de pantalla"
    ],
    # Mujer gestante y lactante
    "maternidad": [
        "mujer gestante","maternidad","control prenatal","licencia de maternidad",
        "lactancia","fuero de maternidad","sala de lactancia"
    ],
}

# ------------------------------------------------------------
# Utilidades de BD (SQLite)
# ------------------------------------------------------------

def _connect(db_path: str) -> sqlite3.Connection:
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn

def _now() -> str:
    return datetime.now().isoformat(timespec="seconds")

def _init_db(db_path: str):
    os.makedirs(os.path.dirname(db_path), exist_ok=True)
    conn = _connect(db_path)
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS cases (
            id TEXT PRIMARY KEY,
            title TEXT,
            status TEXT,
            created_at TEXT,
            updated_at TEXT
        )""")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS parties (
            id TEXT PRIMARY KEY,
            case_id TEXT,
            role TEXT, -- 'accionante' | 'accionado'
            nombres TEXT,
            apellidos TEXT,
            tipo_id TEXT,
            numero_id TEXT,
            email TEXT,
            telefono TEXT,
            direccion TEXT,
            created_at TEXT,
            updated_at TEXT,
            FOREIGN KEY(case_id) REFERENCES cases(id)
        )""")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS sections (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            case_id TEXT,
            name TEXT,
            user_text TEXT,
            ai_text TEXT,
            final_text TEXT,
            needs_llm INTEGER,
            send_order INTEGER,
            status TEXT, -- 'empty'|'draft'|'ai_suggested'|'approved'
            citations_json TEXT,
            updated_at TEXT,
            UNIQUE(case_id, name),
            FOREIGN KEY(case_id) REFERENCES cases(id)
        )""")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS rights_detected (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            case_id TEXT,
            right_name TEXT,
            argument_ai TEXT,
            sources_json TEXT,
            updated_at TEXT,
            UNIQUE(case_id, right_name),
            FOREIGN KEY(case_id) REFERENCES cases(id)
        )""")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS versions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            case_id TEXT,
            section_name TEXT,
            final_text_snapshot TEXT,
            created_at TEXT,
            FOREIGN KEY(case_id) REFERENCES cases(id)
        )""")
    conn.commit()
    conn.close()

def _ensure_sections_for_case(conn: sqlite3.Connection, case_id: str):
    """Crea filas en 'sections' para cada sección del config si no existen."""
    cur = conn.cursor()
    existing = set(
        row["name"] for row in cur.execute(
            "SELECT name FROM sections WHERE case_id=?", (case_id,))
    )
    for name, meta in SECTIONS_CONFIG.items():
        if name in existing:
            continue
        cur.execute("""
            INSERT INTO sections
            (case_id, name, user_text, ai_text, final_text, needs_llm, send_order, status, citations_json, updated_at)
            VALUES (?,?,?,?,?,?,?,?,?,?)
        """, (
            case_id, name, "", "", "",
            int(bool(meta["needs_llm"])) ,
            meta["send_order"] if meta["send_order"] is not None else None,
            "empty", "[]", _now()
        ))
    conn.commit()

def _get_case_bundle(conn: sqlite3.Connection, case_id: str) -> Dict[str, Any]:
    cur = conn.cursor()
    c = cur.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone()
    if not c:
        raise HTTPException(status_code=404, detail="Caso no encontrado")

    parties = [ dict(row) for row in cur.execute(
        "SELECT * FROM parties WHERE case_id=? ORDER BY role, created_at", (case_id,)
    ).fetchall() ]

    sections = [ dict(row) for row in cur.execute(
        "SELECT * FROM sections WHERE case_id=? ORDER BY id", (case_id,)
    ).fetchall() ]

    rights = [ dict(row) for row in cur.execute(
        "SELECT * FROM rights_detected WHERE case_id=? ORDER BY right_name", (case_id,)
    ).fetchall() ]

    return {
        "case": dict(c),
        "parties": parties,
        "sections": sections,
        "rights_detected": rights,
    }

def _upsert_party(conn: sqlite3.Connection, case_id: str, data: Dict[str, Any]) -> Dict[str, Any]:
    """Inserta o actualiza una 'party' por id (si viene) o crea nueva."""
    cur = conn.cursor()
    pid = (data.get("id") or "").strip()
    now = _now()
    fields = ["role", "nombres", "apellidos", "tipo_id", "numero_id", "email", "telefono", "direccion"]

    if pid:
        row = cur.execute("SELECT * FROM parties WHERE id=? AND case_id=?", (pid, case_id)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Persona no encontrada en el caso")
        updates = [f"{f}=: {f}".replace(": ", ":") for f in fields]  # role=:role, nombres=:nombres, ...
        cur.execute(f"""
            UPDATE parties SET {", ".join(updates)}, updated_at=:updated_at
            WHERE id=:id AND case_id=:case_id
        """, {**{f: data.get(f, row[f]) for f in fields},
              "updated_at": now, "id": pid, "case_id": case_id})
        conn.commit()
        out = dict(cur.execute("SELECT * FROM parties WHERE id=?", (pid,)).fetchone())
    else:
        # crear
        new_id = uuid.uuid4().hex[:12]
        cur.execute("""
            INSERT INTO parties
            (id, case_id, role, nombres, apellidos, tipo_id, numero_id, email, telefono, direccion, created_at, updated_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
        """, (
            new_id, case_id,
            data.get("role","").strip(),
            data.get("nombres","").strip(),
            data.get("apellidos","").strip(),
            data.get("tipo_id","").strip(),
            data.get("numero_id","").strip(),
            data.get("email","").strip(),
            data.get("telefono","").strip(),
            data.get("direccion","").strip(),
            now, now
        ))
        conn.commit()
        out = dict(cur.execute("SELECT * FROM parties WHERE id=?", (new_id,)).fetchone())

    # refrescar INTRO con personas reales y autogenerar Notificaciones/Firmas
    _refresh_intro_after_party(conn, case_id)
    _refresh_notifications(conn, case_id)
    _refresh_firmas(conn, case_id)
    return out

def _save_section_user_text(conn: sqlite3.Connection, case_id: str, name: str, text: str) -> Dict[str, Any]:
    # Validación especial de pretensiones (prohibidas pretensiones económicas)
    if name == "pretensiones":
        if _contains_economic_claim(text or ""):
            raise HTTPException(
                status_code=400,
                detail="Las pretensiones económicas están prohibidas en la tutela. Reformula sin montos ni pagos."
            )

    cur = conn.cursor()
    meta = SECTIONS_CONFIG.get(name)
    if not meta:
        raise HTTPException(status_code=404, detail=f"Sección desconocida: {name}")

    res = cur.execute("""
        UPDATE sections SET user_text=?, status=?, updated_at=?
        WHERE case_id=? AND name=?
    """, (text, "draft" if (text or "").strip() else "empty", _now(), case_id, name))
    if res.rowcount == 0:
        raise HTTPException(status_code=404, detail="Sección no existe para este caso")
    conn.commit()

    return dict(cur.execute("""
        SELECT * FROM sections WHERE case_id=? AND name=?
    """, (case_id, name)).fetchone())

def _save_section_ai(conn: sqlite3.Connection, case_id: str, name: str,
                     ai_text: str, citations: Optional[List[Dict[str, Any]]] = None) -> Dict[str, Any]:
    # También filtramos pretensiones si por error se invoca IA allí
    if name == "pretensiones" and _contains_economic_claim(ai_text or ""):
        raise HTTPException(
            status_code=400,
            detail="La IA propuso pretensiones económicas, lo cual está prohibido. Edita el texto y quita montos."
        )

    cur = conn.cursor()
    meta = SECTIONS_CONFIG.get(name)
    if not meta:
        raise HTTPException(status_code=404, detail=f"Sección desconocida: {name}")

    citations_json = json.dumps(citations or [], ensure_ascii=False)
    res = cur.execute("""
        UPDATE sections SET ai_text=?, status=?, citations_json=?, updated_at=?
        WHERE case_id=? AND name=?
    """, (ai_text, "ai_suggested" if (ai_text or "").strip() else "draft",
          citations_json, _now(), case_id, name))
    if res.rowcount == 0:
        raise HTTPException(status_code=404, detail="Sección no existe para este caso")
    conn.commit()

    return dict(cur.execute("SELECT * FROM sections WHERE case_id=? AND name=?",
                            (case_id, name)).fetchone())

def _approve_section(conn: sqlite3.Connection, case_id: str, name: str, source: str = "ai") -> Dict[str, Any]:
    cur = conn.cursor()
    row = cur.execute("SELECT * FROM sections WHERE case_id=? AND name=?", (case_id, name)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Sección no existe para este caso")

    final_text = (row["ai_text"] if source == "ai" else row["user_text"]) or ""
    # Validación de pretensiones antes de aprobar
    if name == "pretensiones" and _contains_economic_claim(final_text):
        raise HTTPException(status_code=400, detail="Pretensiones económicas prohibidas. Edita y elimina montos.")

    cur.execute("""
        UPDATE sections SET final_text=?, status=?, updated_at=?
        WHERE case_id=? AND name=?
    """, (final_text, "approved" if final_text.strip() else row["status"],
          _now(), case_id, name))

    # Guarda versión
    cur.execute("""
        INSERT INTO versions (case_id, section_name, final_text_snapshot, created_at)
        VALUES (?,?,?,?)
    """, (case_id, name, final_text, _now()))
    conn.commit()

    return dict(cur.execute("SELECT * FROM sections WHERE case_id=? AND name=?",
                            (case_id, name)).fetchone())

def _invalidate_sections(conn: sqlite3.Connection, case_id: str, names: List[str]):
    """Limpia ai_text/final_text y baja estado a 'draft' para forzar regeneración posterior."""
    if not names:
        return
    cur = conn.cursor()
    now = _now()
    q_marks = ",".join(["?"] * len(names))
    cur.execute(f"""
        UPDATE sections
        SET ai_text='', final_text='', status='draft', updated_at=?
        WHERE case_id=? AND name IN ({q_marks})
    """, [now, case_id, *names])
    conn.commit()

def _fold(s: str) -> str:
    # lower + quita acentos/diacríticos
    s = (s or "").lower()
    return "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c))

# Precompilamos patrones una sola vez para rendimiento y precisión
RIGHT_PATTERNS: Dict[str, List[re.Pattern]] = {}
for right, kws in RIGHTS_LEXICON.items():
    pats = []
    for kw in kws:
        k = _fold(kw).strip()
        if not k:
            continue
        # escapamos y permitimos espacios flexibles
        k_escaped = re.escape(k)
        k_escaped = re.sub(r"\\\s+", r"\\s+", k_escaped)  # por si en el futuro hay \s en keywords
        k_escaped = re.sub(r"\s+", r"\\s+", k_escaped)    # "derecho de peticion" -> "derecho\s+de\s+peticion"
        # si es token alfanumérico (sin espacios/puntuación), aplicamos bordes de palabra
        if re.fullmatch(r"[a-z0-9]+", k):
            pat = re.compile(rf"\b{k_escaped}\b")
        else:
            pat = re.compile(k_escaped)
        pats.append(pat)
    RIGHT_PATTERNS[right] = pats

def _detect_rights(text: str) -> List[str]:
    """
    Devuelve la lista ORDENADA de derechos detectados usando exactamente el nombre
    de la clave de RIGHTS_LEXICON (p. ej., 'salud', 'vida digna').
    """
    if not text:
        return []
    t = _fold(text)
    found = set()
    for right, patterns in RIGHT_PATTERNS.items():
        if any(p.search(t) for p in patterns):
            found.add(right)
    return sorted(found)

def _set_right(conn: sqlite3.Connection, case_id: str, right_name: str,
               argument_ai: str = "", sources: Optional[List[Dict[str, Any]]] = None) -> Dict[str, Any]:
    cur = conn.cursor()
    now = _now()
    js = json.dumps(sources or [], ensure_ascii=False)
    cur.execute("""
        INSERT INTO rights_detected (case_id, right_name, argument_ai, sources_json, updated_at)
        VALUES (?,?,?,?,?)
        ON CONFLICT(case_id, right_name) DO UPDATE SET
            argument_ai=excluded.argument_ai,
            sources_json=excluded.sources_json,
            updated_at=excluded.updated_at
    """, (case_id, right_name, argument_ai, js, now))
    conn.commit()
    return dict(cur.execute("""
        SELECT * FROM rights_detected WHERE case_id=? AND right_name=?
    """, (case_id, right_name)).fetchone())

# -------- Helpers de personas / intro / notificaciones / firmas --------
# # Mejora del BLOQUE 4:
# - Soporta variaciones: "contra"/"en contra de"
# - Reemplaza placeholders tipo "XXX..." y frases genericas "los accionados"/"el accionado"
# - Evita duplicar sustituciones y normaliza espacios
X_RE = re.compile(r"X{3,}", flags=re.I)  # Marcadores tipo 'XXXX...'
RE_SENTINEL = re.compile(
    r"\b(l[oa]s?\s+accionad[oa]s?)\b",  # 'los accionados', 'el accionado', 'la accionada', etc.
    flags=re.I
)
RE_CONTRA_BLOCK = re.compile(
    r"(ACCI[ÓO]N\s+DE\s+TUTELA\s+(?:en\s+)?contra\s+)(.*?)(,\s+con\s+el\s+objeto)",
    flags=re.I | re.S
)

def _normalize_spaces(s: str) -> str:
    return re.sub(r"[ \t]+", " ", (s or "").strip())

def _replace_contra_segment(txt: str, accionados_inline: str) -> str:
    """
    Reemplaza el tramo '... TUTELA (en) contra <AQUI> , con el objeto ...' por los accionados reales.
    Si no encuentra el patrón, se limita a devolver el texto original.
    """
    if not (txt or "").strip() or not (accionados_inline or "").strip():
        return txt or ""

    def _sub(m):
        pre = m.group(1)
        mid = accionados_inline
        post = m.group(3)
        return f"{pre}{mid}{post}"

    return RE_CONTRA_BLOCK.sub(_sub, txt)

def _compose_people_inline(conn, case_id: str, role: str) -> str:
    cur = conn.cursor()
    rows = cur.execute(
        "SELECT nombres, apellidos, tipo_id, numero_id FROM parties "
        "WHERE case_id=? AND role=? ORDER BY created_at", (case_id, role)
    ).fetchall()

    out = []
    for r in rows:
        nombre = " ".join([(r["nombres"] or "").strip(), (r["apellidos"] or "").strip()]).strip()
        ident  = " ".join([(r["tipo_id"] or "").strip(), (r["numero_id"] or "").strip()]).strip()
        if nombre and ident:
            out.append(f"{nombre} — {ident}")
        elif nombre:
            out.append(nombre)
    return ", ".join(out)

def _refresh_intro_after_party(conn, case_id: str) -> None:
    """
    Regenera o actualiza la INTRO con los accionantes/accionados reales.
    - Si la intro está vacía → crea un texto base con las partes.
    - Si existe → sustituye placeholders y el bloque 'contra <...> , con el objeto'.
    """
    cur = conn.cursor()
    row = cur.execute(
        "SELECT user_text, ai_text, final_text FROM sections WHERE case_id=? AND name='intro'",
        (case_id,)
    ).fetchone()
    current = _normalize_spaces((row and (row["user_text"] or row["final_text"] or row["ai_text"])) or "")

    acc_str = _compose_people_inline(conn, case_id, "accionante")
    ads_str = _compose_people_inline(conn, case_id, "accionado")

    if not current:
        accionado_display  = ads_str or "los accionados"
        accionante_display = acc_str or "El accionante"
        intro_text = (
            f"{accionante_display}, identificado como aparece al pie de mi firma, actuando en nombre propio, "
            f"invocando el artículo 86 de la Constitución Política, acudo ante su Despacho para instaurar "
            f"ACCIÓN DE TUTELA contra {accionado_display}, con el objeto de que se protejan los derechos "
            f"constitucionales fundamentales que a continuación enuncio y los cuales se fundamentan en los siguientes hechos."
        )
        _save_section_user_text(conn, case_id, "intro", intro_text)
        return

    new_txt = current
    if ads_str:
        # 1) Sustituye placeholders tipo 'XXX...'
        new_txt = X_RE.sub(ads_str, new_txt)
        # 2) Sustituye frases genéricas 'los/el/la accionado(s)'
        new_txt = RE_SENTINEL.sub(ads_str, new_txt)
        # 3) Sustituye tramo '... contra <...> , con el objeto'
        new_txt = _replace_contra_segment(new_txt, ads_str)

    new_txt = _normalize_spaces(new_txt)
    if new_txt != current:
        _save_section_user_text(conn, case_id, "intro", new_txt)

def _compose_notifications_text(conn, case_id: str) -> str:
    """Genera el texto de Notificaciones para accionantes y accionados desde parties."""
    cur = conn.cursor()

    def fmt_block(role_label: str, role_key: str) -> str:
        people = cur.execute(
            "SELECT * FROM parties WHERE case_id=? AND role=? ORDER BY created_at",
            (case_id, role_key)
        ).fetchall()
        if not people:
            return f"{role_label}:\n(sin datos aún)"
        lines = []
        for p in people:
            nombre = " ".join([(p["nombres"] or "").strip(), (p["apellidos"] or "").strip()]).strip()
            campos = []
            if p["telefono"]:
                campos.append(f"Tel: {p['telefono'].strip()}")
            if p["email"]:
                campos.append(f"Correo: {p['email'].strip()}")
            if p["direccion"]:
                campos.append(f"Dirección: {p['direccion'].strip()}")
            tail = " | ".join(campos) if campos else "(sin datos de contacto)"
            lines.append(f"Se notificará a {nombre} — {tail}".strip())
        return f"{role_label}:\n" + "\n".join(lines)

    acc = fmt_block("Accionante(s)", "accionante")
    ads = fmt_block("Accionado(s)", "accionado")
    return acc + "\n\n" + ads

def _compose_firmas_text(conn, case_id: str) -> str:
    """Genera el bloque de firmas con nombres + identificación de todos los accionantes."""
    cur = conn.cursor()
    rows = cur.execute(
        "SELECT * FROM parties WHERE case_id=? AND role='accionante' ORDER BY created_at",
        (case_id,)
    ).fetchall()
    if not rows:
        return "(agrega al menos un accionante para firmar)"
    lines = []
    for r in rows:
        nombre = " ".join([(r["nombres"] or "").strip(), (r["apellidos"] or "").strip()]).strip() or "(sin nombre)"
        ident  = " ".join([(r["tipo_id"] or "").strip(), (r["numero_id"] or "").strip()]).strip() or "ID — (pendiente)"
        lines.append(f"{nombre} — {ident}")
    return "\n".join(lines)

def _refresh_notifications(conn, case_id: str) -> None:
    txt = _compose_notifications_text(conn, case_id)
    _save_section_user_text(conn, case_id, "notificaciones", txt)

def _refresh_firmas(conn, case_id: str) -> None:
    txt = _compose_firmas_text(conn, case_id)
    _save_section_user_text(conn, case_id, "firmas", txt)

# ------------------------------------------------------------
# LLM / RAG helpers
# ------------------------------------------------------------

def _docs_for_prompt(retriever, query: str, k: int = 5) -> Tuple[List[str], List[Dict[str, Any]]]:
    """
    Obtiene documentos del retriever usando la API moderna (.invoke) y, si no existe,
    hace fallback a .get_relevant_documents. Devuelve (chunks_para_prompt, citas_struct).
    """
    if not retriever:
        return [], []

    def _normalize_docs(docs_obj):
        if docs_obj is None:
            return []
        if isinstance(docs_obj, dict):
            for key in ("documents", "docs", "output"):
                val = docs_obj.get(key)
                if isinstance(val, list):
                    return val
            return []
        if isinstance(docs_obj, list):
            return docs_obj
        return [docs_obj]

    docs = []
    try:
        docs = _normalize_docs(retriever.invoke(query))
    except AttributeError:
        try:
            docs = _normalize_docs(retriever.get_relevant_documents(query))
        except Exception:
            docs = []
    except Exception:
        try:
            docs = _normalize_docs(retriever.get_relevant_documents(query))
        except Exception:
            docs = []

    chunks, cites = [], []
    for d in docs[:k]:
        content = getattr(d, "page_content", None)
        meta = getattr(d, "metadata", None)
        if content is None and isinstance(d, dict):
            content = d.get("page_content") or d.get("content") or ""
            meta = d.get("metadata") or {}
        content = (content or "").strip()
        meta = meta or {}
        title = meta.get("title") or meta.get("source") or "doc"
        if content:
            snippet = content[:600]
            chunks.append(f"- {title}: {snippet}")
            cites.append({"title": title, "snippet": snippet, "meta": meta})
    return chunks, cites

def _generate_fundamentos_juridicos(llm, ctx: Dict[str, Any]) -> str:
    """
    Genera FUNDAMENTOS JURÍDICOS en 4 sub-llamadas:
    1) Procedencia, 2) Problema jurídico, 3) Reglas (jurisprudenciales/legales), 4) Caso concreto.
    Devuelve un único texto ya ordenado (1..4).
    """
    if not llm:  # Fallback simple si no hay LLM
        return (
            "1) Procedencia: analiza subsidiariedad, inmediatez, legitimación y perjuicio irremediable.\n\n"
            "2) Problema jurídico: formula una pregunta clara según los hechos.\n\n"
            "3) Reglas: enuncia reglas jurisprudenciales y legales pertinentes de forma sintética.\n\n"
            "4) Caso concreto: subsume los hechos a las reglas y explica la vulneración."
        )

    H = ctx.get("hechos", "")
    D = ctx.get("derechos_vulnerados", "") or "\n".join(ctx.get("derechos_detectados_dic", []))
    P = ctx.get("pruebas", "")

    prompts = {
        "procedencia": (
            "Redacta el apartado 'Procedencia' de una acción de tutela (Colombia) en 1–3 párrafos, "
            "analizando subsidiariedad, inmediatez, legitimación por activa/pasiva y perjuicio irremediable. "
            "Ancla a hechos (H#) si corresponde. Sin metadiscurso.\n\n"
            f"Hechos (H#):\n{H}\n\nDerechos (D#):\n{D}\n\nPruebas (P#):\n{P}"
        ),
        "problema": (
            "Formula el 'Problema jurídico' como UNA pregunta clara y completa, en una sola línea, "
            "derivada de los hechos y los derechos invocados. Sin explicaciones.\n\n"
            f"Hechos (H#):\n{H}\n\nDerechos (D#):\n{D}"
        ),
        "reglas": (
            "Enuncia 'Reglas jurisprudenciales y legales' en 3–6 ítems breves (sin citas extensas). "
            "Cada ítem: regla clara aplicable al caso (enunciado general). Sin inventar números/fechas.\n\n"
            f"Hechos (H#):\n{H}\n\nDerechos (D#):\n{D}"
        ),
        "caso": (
            "Redacta 'Caso concreto' en 2–4 párrafos breves: subsume hechos a reglas, "
            "explica por qué se configura (o no) la vulneración. Referencia (H#) o (P#) cuando proceda. "
            "Sin frases de cierre grandilocuentes.\n\n"
            f"Hechos (H#):\n{H}\n\nDerechos (D#):\n{D}\n\nPruebas (P#):\n{P}"
        ),
    }

    out = {}
    for key, pr in prompts.items():
        try:
            r = llm.invoke(pr)
            out[key] = (getattr(r, "content", None) or str(r) or "").strip()
        except Exception:
            out[key] = ""

    texto = (
        "1) Procedencia:\n" + (out["procedencia"] or "") + "\n\n"
        "2) Problema jurídico:\n" + (out["problema"] or "") + "\n\n"
        "3) Reglas jurisprudenciales y legales:\n" + (out["reglas"] or "") + "\n\n"
        "4) Caso concreto:\n" + (out["caso"] or "")
    ).strip()
    return texto

def _get_best_text(row: Optional[sqlite3.Row]) -> str:
    if not row:
        return ""
    return (row["final_text"] or row["ai_text"] or row["user_text"] or "").strip()

def _llm_improve_for_section(
    name: str,
    user_text: str,
    ctx: Dict[str, Any],
    llm=None,
    retriever=None
) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Devuelve (ai_text, citations_list). Prohibido metadiscurso; entrega el contenido final.
    Ajustes:
    - DERECHOS_VULNERADOS: añade derechos detectados por diccionario al prompt.
    - FUNDAMENTOS_JURÍDICOS: 4 sub-llamadas (procedencia, problema, reglas, caso).
    - FUNDAMENTOS_DE_DERECHO (RAG): lista numerada basada en los docs recuperados.
    """
    # Fallback si no hay LLM
    if not llm:
        base = (user_text or "").strip()
        if not base:
            if name == "derechos_vulnerados":
                det = ctx.get("derechos_detectados_dic", [])
                base = ("Derechos detectados: " + ", ".join(det) + "\n\n" if det else "") + \
                       (ctx.get("hechos","") or "")
            elif name == "fundamentos_juridicos":
                return _generate_fundamentos_juridicos(None, ctx), []
            elif name == "fundamentos_de_derecho":
                return "1) C.P., art. 86\n2) D. 2591 de 1991, arts. 5, 6 y 42", []
            elif name == "ref":
                prompt_parts += [
                    guides.get(name, ""),
                    f"Derechos:\n{ctx.get('derechos_vulnerados','')}\n\n"
                    f"Fundamentos Jurídicos:\n{ctx.get('fundamentos_juridicos','')}\n\n"
                    f"Fundamentos de Derecho:\n{ctx.get('fundamentos_de_derecho','')}\n\n"
                    "Partes (úsalas tal cual, sin corchetes):\n"
                    f"Accionantes: {ctx.get('accionantes_inline','') or '(sin registrar)'}\n"
                    f"Accionados: {ctx.get('accionados_inline','') or '(sin registrar)'}"
                ]
            elif name == "pruebas_y_anexos":
                base = "Normaliza el listado de pruebas/anexos (uno por línea)."
            elif name == "pretensiones":
                base = "Organiza pretensiones en lista numerada clara y ejecutable (sin economía)."
        return base or user_text or "", []

    # ---- RAG sólo para Fundamentos de Derecho ----
    rag_chunks, citations = [], []
    if name == "fundamentos_de_derecho" and retriever:
        query = ctx.get("fundamentos_juridicos") or ctx.get("hechos") or ""
        rag_chunks, citations = _docs_for_prompt(retriever, query, k=8)

    # ---- Enrutamiento por sección ----
    if name == "fundamentos_juridicos":
        return _generate_fundamentos_juridicos(llm, ctx), []

    prompt_parts = ["Eres un redactor jurídico colombiano especializado en acciones de tutela."]

    if name == "derechos_vulnerados":
        det = ctx.get("derechos_detectados_dic", [])
        det_line = "Derechos detectados (diccionario): " + (", ".join(det) if det else "(ninguno)")
        prompt_parts += [
            guides.get(name, ""),
            f"Hechos (H#):\n{ctx.get('hechos','')}",
            det_line
        ]
    elif name == "fundamentos_de_derecho":
        # Debe basarse en RAG: no inventes
        prompt_parts += [
            guides.get(name, ""),
            "Base EXCLUSIVA para citar (no inventes, usa lo siguiente):\n" + ("\n".join(rag_chunks) if rag_chunks else "-"),
            "ENTREGA SOLO una lista numerada de normas/sentencias reales, formateadas como en la guía. Sin comentarios.",
        ]
    elif name == "ref":
        prompt_parts += [
            guides.get(name, ""),
            f"Derechos:\n{ctx.get('derechos_vulnerados','')}\n\n"
            f"Fundamentos Jurídicos:\n{ctx.get('fundamentos_juridicos','')}\n\n"
            f"Fundamentos de Derecho:\n{ctx.get('fundamentos_de_derecho','')}"
        ]
    else:
        # genérico
        prompt_parts += [
            guides.get(name, "Mejora la redacción sin inventar."),
            f"Contexto:\n{json.dumps(ctx, ensure_ascii=False)[:1500]}"
        ]

    prompt_parts.append("ENTREGA SOLO EL CONTENIDO SOLICITADO, sin introducciones.")
    if (user_text or "").strip():
        prompt_parts.append(f"Texto del usuario (si aplica):\n\"\"\"\n{user_text.strip()}\n\"\"\"")

    try:
        resp = llm.invoke("\n\n".join(prompt_parts))
        content = getattr(resp, "content", None) or str(resp)
        return content.strip(), citations
    except Exception:
        return (user_text or ctx.get("hechos","") or "").strip(), []

# -------- Prompts específicos (sugerencias y cadena) --------
PROMPT_SUGIERE_PRETENSIONES = (
    "Con base en los HECHOS (depurados) y las pretensiones del usuario ya redactadas, "
    "sugiere 3–5 pretensiones adicionales legítimas, razonables y NO económicas. "
    "Devuelve SOLO una lista numerada corta, sin preámbulos.\n\n"
    "HECHOS:\n{hechos}\n\nPRETENSIONES (usuario/limpiadas):\n{pret}\n"
)

PROMPT_CADENA_JSON = (
    "A partir de los HECHOS (depurados) y las PRETENSIONES (mejoradas), genera para una acción de tutela:\n"
    "1) DERECHOS VULNERADOS (texto)\n"
    "2) FUNDAMENTOS JURÍDICOS (texto)\n"
    "3) FUNDAMENTOS DE DERECHO (lista numerada de normas/sentencias reales; no expliques)\n"
    "4) REF (línea de referencia)\n\n"
    "Devuelve exclusivamente este JSON MINIMAL (sin backticks ni explicación):\n"
    "{\n"
    ' "derechos_vulnerados": "...",\n'
    ' "fundamentos_juridicos": "...",\n'
    ' "fundamentos_de_derecho": "...",\n'
    ' "ref": "..."\n'
    "}\n\n"
    "HECHOS:\n{hechos}\n\nPRETENSIONES:\n{pret}\n"
)

# ------------------------------------------------------------
# Export a Word (.docx) y composición de texto
# ------------------------------------------------------------

def _pick_final(row: sqlite3.Row) -> str:
    return (row["final_text"] or row["ai_text"] or row["user_text"] or "").strip()

def _join_people(rows: List[sqlite3.Row]) -> str:
    out = []
    for r in rows:
        nombre = " ".join([r["nombres"] or "", r["apellidos"] or ""]).strip()
        ident  = " ".join([r["tipo_id"] or "", r["numero_id"] or ""]).strip()
        label = nombre if not ident else f"{nombre} — {ident}"
        out.append(label)
    return "; ".join([x for x in out if x])

def _compose_full_text(conn: sqlite3.Connection, case_id: str) -> str:
    """Devuelve el documento completo en texto plano, para previsualizar en UI."""
    cur = conn.cursor()

    sections = {row["name"]: row for row in cur.execute(
        "SELECT * FROM sections WHERE case_id=?", (case_id,)
    ).fetchall()}

    accionantes = [row for row in cur.execute(
        "SELECT * FROM parties WHERE case_id=? AND role='accionante'", (case_id,)
    ).fetchall()]
    accionados = [row for row in cur.execute(
        "SELECT * FROM parties WHERE case_id=? AND role='accionado'", (case_id,)
    ).fetchall()]

    ref     = _pick_final(sections.get("ref")) if sections.get("ref") else ""
    intro   = _pick_final(sections.get("intro")) if sections.get("intro") else ""
    hechos  = _pick_final(sections.get("hechos")) if sections.get("hechos") else ""
    der_vuln= _pick_final(sections.get("derechos_vulnerados")) if sections.get("derechos_vulnerados") else ""
    fund_j  = _pick_final(sections.get("fundamentos_juridicos")) if sections.get("fundamentos_juridicos") else ""
    fund_d  = _pick_final(sections.get("fundamentos_de_derecho")) if sections.get("fundamentos_de_derecho") else ""
    pret    = _pick_final(sections.get("pretensiones")) if sections.get("pretensiones") else ""
    notifs  = _pick_final(sections.get("notificaciones")) if sections.get("notificaciones") else ""
    firmas  = _pick_final(sections.get("firmas")) if sections.get("firmas") else ""
    juramento = _pick_final(sections.get("cumplimiento_art_37")) if sections.get("cumplimiento_art_37") else ""

    # pruebas/anexos (nuevo y compatibilidad)
    pya        = _pick_final(sections.get("pruebas_y_anexos")) if sections.get("pruebas_y_anexos") else ""
    pruebas_old= _pick_final(sections.get("pruebas")) if sections.get("pruebas") else ""
    anexos_old = _pick_final(sections.get("anexos")) if sections.get("anexos") else ""
    pa_text    = pya or "\n".join([t for t in [pruebas_old, anexos_old] if t.strip()])

    parts = []
    parts.append(HEADER_FIXED)
    if ref:
        parts.append("\nREF: " + ref)

    parts.append("\nACCIONANTE(S): " + (_join_people(accionantes) or ""))
    parts.append("ACCIONADO(S): " + (_join_people(accionados) or ""))

    if intro:
        parts.append("\n## Introducción")
        parts.append(intro)

    if hechos:
        parts.append("\n## Hechos")
        parts.append(hechos)

    if der_vuln:
        parts.append("\n## Derechos vulnerados")
        parts.append(der_vuln)

    if fund_j:
        parts.append("\n## Fundamentos jurídicos")
        parts.append(fund_j)

    parts.append("\n## Pruebas y Anexos")
    if pa_text.strip():
        parts.append(
            "Con el fin de establecer la vulneración de los derechos, solicito señor Juez se sirva "
            "tener en cuenta las siguientes pruebas y anexos:"
        )
        for i, line in enumerate([l for l in pa_text.splitlines() if l.strip()], start=1):
            parts.append(f"{i}. {line.strip()}")
    else:
        parts.append("(sin registros)")

    if pret.strip():
        parts.append("\n## Pretensiones")
        for i, line in enumerate([l for l in pret.splitlines() if l.strip()], start=1):
            parts.append(f"{i}. {line.strip()}")

    if fund_d:
        parts.append("\n## Fundamentos de derecho")
        for i, line in enumerate([l for l in fund_d.splitlines() if l.strip()], start=1):
            parts.append(f"{i}. {line.strip()}")

    if juramento.strip():
        parts.append("\n## Cumplimiento art. 37 del Decreto 2591/1991 — Juramento")
        parts.append(juramento.strip())

    if notifs.strip():
        parts.append("\n## Notificaciones")
        parts.append(notifs.strip())

    firmas_txt = firmas.strip() or _compose_firmas_text(conn, case_id)
    if firmas_txt:
        parts.append("\n## Firmas")
        parts.append(firmas_txt)

    return "\n".join(parts).strip()

def _export_docx(conn: sqlite3.Connection, case_id: str, export_dir: str) -> Dict[str, str]:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:
        raise RuntimeError("Instala 'python-docx' para exportar a Word: pip install python-docx") from e

    # -------- Helpers de estilo --------
    FONT_NAME = "Times New Roman"
    FONT_SIZE = 12

    def _set_normal_style(doc):
        st = doc.styles["Normal"]
        st.font.name = FONT_NAME
        st.font.size = Pt(FONT_SIZE)

    def _p(doc, text, *, bold=False, upper=False, align=None):
        if upper:
            text = (text or "").upper()
        par = doc.add_paragraph()
        if align is not None:
            par.alignment = align
        run = par.add_run(text or "")
        run.bold = bool(bold)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        return par

    def _title(doc, text):
        # Título en MAYÚSCULAS, centrado y en NEGRITA (sin estilos de Heading para unificar fuente)
        return _p(doc, text, bold=True, upper=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    def _body(doc, text):
        par = _p(doc, text or "", bold=False, upper=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        return par

    def _join_people(rows: List[sqlite3.Row]) -> str:
        out = []
        for r in rows:
            nombre = " ".join([(r["nombres"] or "").strip(), (r["apellidos"] or "").strip()]).strip()
            ident  = " ".join([(r["tipo_id"] or "").strip(), (r["numero_id"] or "").strip()]).strip()
            label = nombre if not ident else f"{nombre} — {ident}"
            if label:
                out.append(label)
        return "; ".join(out)

    cur = conn.cursor()
    case = cur.execute("SELECT * FROM cases WHERE id=?", (case_id,)).fetchone()
    if not case:
        raise HTTPException(status_code=404, detail="Caso no encontrado")

    sections = {row["name"]: row for row in cur.execute(
        "SELECT * FROM sections WHERE case_id=?", (case_id,)
    ).fetchall()}

    accionantes = [row for row in cur.execute(
        "SELECT * FROM parties WHERE case_id=? AND role='accionante'", (case_id,)
    ).fetchall()]
    accionados = [row for row in cur.execute(
        "SELECT * FROM parties WHERE case_id=? AND role='accionado'", (case_id,)
    ).fetchall()]

    # Selección de textos
    def pick(name):
        r = sections.get(name)
        return (r["final_text"] or r["ai_text"] or r["user_text"] or "").strip() if r else ""

    ref        = pick("ref")
    intro      = pick("intro")
    hechos     = pick("hechos")
    der_vuln   = pick("derechos_vulnerados")
    fund_j     = pick("fundamentos_juridicos")
    fund_d     = pick("fundamentos_de_derecho")
    pret       = pick("pretensiones")
    notifs     = pick("notificaciones")
    juramento  = pick("cumplimiento_art_37") or "JURAMENTO: Manifiesto bajo la gravedad del juramento que no se ha presentado ninguna otra acción de tutela por los mismos hechos y derechos."
    pya        = pick("pruebas_y_anexos") or "\n".join([pick("pruebas"), pick("anexos")]).strip()

    # ---- Crear doc
    os.makedirs(export_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    basename  = f"tutela_{case_id}_{ts}"
    docx_path = os.path.join(export_dir, f"{basename}.docx")
    json_path = os.path.join(export_dir, f"{basename}.json")

    doc = Document()
    _set_normal_style(doc)

    # ===== Encabezado (SIN la línea extra), MAYÚSCULA + NEGRITA =====
    for line in HEADER_FIXED.split("\n"):
        _p(doc, line, bold=True, upper=True)

    # REF
    if ref:
        _p(doc, "", bold=False)  # espacio
        par = doc.add_paragraph()
        run1 = par.add_run("REF: ")
        run1.bold = True; run1.font.name = FONT_NAME; run1.font.size = Pt(FONT_SIZE)
        run2 = par.add_run(ref)
        run2.font.name = FONT_NAME; run2.font.size = Pt(FONT_SIZE)

    # Partes (nombres e identificaciones en MAYÚSCULA + NEGRITA)
    _p(doc, "", bold=False)  # espacio
    par = doc.add_paragraph()
    r1 = par.add_run("ACCIONANTE(S): ")
    r1.bold = True; r1.font.name = FONT_NAME; r1.font.size = Pt(FONT_SIZE)
    r2 = par.add_run(_join_people(accionantes).upper())
    r2.bold = True; r2.font.name = FONT_NAME; r2.font.size = Pt(FONT_SIZE)

    par = doc.add_paragraph()
    r1 = par.add_run("ACCIONADO(S): ")
    r1.bold = True; r1.font.name = FONT_NAME; r1.font.size = Pt(FONT_SIZE)
    r2 = par.add_run(_join_people(accionados).upper())
    r2.bold = True; r2.font.name = FONT_NAME; r2.font.size = Pt(FONT_SIZE)

    # ===== Secciones =====
    if intro:
        _title(doc, "Introducción")
        _body(doc, intro)

    if hechos:
        _title(doc, "Hechos")
        _body(doc, hechos)

    if der_vuln:
        _title(doc, "Derechos vulnerados")
        _body(doc, der_vuln)

    if fund_j:
        _title(doc, "Fundamentos jurídicos")
        _body(doc, fund_j)

    # Pruebas y Anexos
    _title(doc, "Pruebas y Anexos")
    _body(doc, "Con el fin de establecer la vulneración de los derechos, solicito señor Juez se sirva tener en cuenta las siguientes pruebas y anexos:")
    if pya.strip():
        for i, line in enumerate([l for l in pya.splitlines() if l.strip()], start=1):
            _body(doc, f"{i}. {line.strip()}")

    # Pretensiones
    if pret.strip():
        _title(doc, "Pretensiones")
        for i, line in enumerate([l for l in pret.splitlines() if l.strip()], start=1):
            _body(doc, f"{i}. {line.strip()}")

    # Fundamentos de derecho
    if fund_d:
        _title(doc, "Fundamentos de derecho")
        for i, line in enumerate([l for l in fund_d.splitlines() if l.strip()], start=1):
            _body(doc, f"{i}. {line.strip()}")

    # Juramento
    if juramento:
        _title(doc, "Cumplimiento art. 37 del Decreto 2591/1991 — Juramento")
        _body(doc, juramento)

    # Notificaciones
    if notifs:
        _title(doc, "Notificaciones")
        _body(doc, notifs)

    # ===== FIRMAS personalizadas =====
    _p(doc, "", bold=False)
    _p(doc, "FIRMAS:", bold=True, upper=True)
    _p(doc, "", bold=False); _p(doc, "", bold=False); _p(doc, "", bold=False)  # 3 saltos

    # Bloques por cada accionante: línea para firma + NOMBRE + TIPO_ID NÚMERO (todo MAYÚSCULA + NEGRITA)
    if not accionantes:
        # Si no hay accionantes, deja un bloque vacío para firmar
        _p(doc, "______________________________", align=WD_ALIGN_PARAGRAPH.CENTER)
        _p(doc, "(NOMBRE DEL ACCIONANTE)", bold=True, upper=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _p(doc, "(TIPO DE ID Y NÚMERO)", bold=True, upper=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        for a in accionantes:
            nombre = " ".join([(a["nombres"] or "").strip(), (a["apellidos"] or "").strip()]).strip() or "(SIN NOMBRE)"
            ident  = " ".join([(a["tipo_id"] or "").strip(), (a["numero_id"] or "").strip()]).strip() or "ID — (PENDIENTE)"
            _p(doc, "______________________________", align=WD_ALIGN_PARAGRAPH.CENTER)
            _p(doc, nombre, bold=True, upper=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _p(doc, ident, bold=True, upper=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _p(doc, "", bold=False)  # espacio entre firmantes

    # Guardar DOCX
    doc.save(docx_path)

    # Export JSON del bundle
    bundle = _get_case_bundle(conn, case_id)
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(bundle, f, ensure_ascii=False, indent=2)

    return {
        "docx_url": f"/exports/{os.path.basename(docx_path)}",
        "json_url": f"/exports/{os.path.basename(json_path)}",
    }

# ------------------------------------------------------------
# Pydantic models
# ------------------------------------------------------------

class CaseCreateResp(BaseModel):
    case_id: str

class CaseListItem(BaseModel):
    case_id: str
    title: str
    status: str
    updated_at: str

class PartyUpsertReq(BaseModel):
    id: Optional[str] = None
    role: str  # 'accionante' | 'accionado'
    nombres: Optional[str] = ""
    apellidos: Optional[str] = ""
    tipo_id: Optional[str] = ""
    numero_id: Optional[str] = ""
    email: Optional[str] = ""
    telefono: Optional[str] = ""
    direccion: Optional[str] = ""

class SectionSaveReq(BaseModel):
    user_text: str

class SectionImproveResp(BaseModel):
    ai_text: str
    citations: List[Dict[str, Any]] = []

class SectionApproveReq(BaseModel):
    source: str = "ai"  # 'ai' | 'user'

class RightsDetectResp(BaseModel):
    rights: List[str]

class RunPipelineResp(BaseModel):
    ran: List[str]

class ExportDocxResp(BaseModel):
    docx_url: str
    json_url: str

class ComposeFinalResp(BaseModel):
    full_text: str

# ------------------------------------------------------------
# Reglas / Validaciones
# ------------------------------------------------------------

# Palabras/expresiones indicativas de pretensiones económicas
ECON_PATTERNS = [
    r"\$",
    r"\bCOP\b",
    r"\bUSD\b",
    r"\beuros?\b",
    r"\bpesos?\b",
    r"\bmillones?\b",
    r"\bmonto\b",
    r"\bcuant[ií]a\b",
    r"\bindemnizaci[oó]n\b",
    r"\bpago[s]?\b",
    r"\bpagar\b",
    r"\bremuneraci[oó]n\b",
    r"\bcompensaci[oó]n\b",
    r"\blucro\b",
    r"\bda[ñn]o[s]?\b\s*(material|inmaterial|moral)",
    r"\binter[eé]s(es)?\b"
]
ECON_RE = re.compile("|".join(ECON_PATTERNS), flags=re.I)

def _contains_economic_claim(text: str) -> bool:
    # Freno desactivado: no bloquear pretensiones económicas
    return False

# ------------------------------------------------------------
# Gating / Dependencias para IA
# ------------------------------------------------------------

def _check_dependencies_or_409(conn: sqlite3.Connection, case_id: str, section_name: str):
    """
    Asegura el orden lógico cuando se llama improve/ensure manualmente.
    - Derechos vulnerados: requiere Hechos mejorados.
    - Fundamentos jurídicos: requiere Derechos vulnerados + Hechos (y Pruebas si existen).
    - Fundamentos de derecho: requiere Fundamentos jurídicos.
    - REF: requiere Derechos + Fund. jurídicos + Fund. de derecho.
    """
    cur = conn.cursor()

    def _ready(name: str) -> bool:
        row = cur.execute("SELECT status, ai_text, final_text, user_text FROM sections WHERE case_id=? AND name=?",
                          (case_id, name)).fetchone()
        if not row:
            return False
        txt = (row["final_text"] or row["ai_text"] or row["user_text"] or "").strip()
        return bool(txt)

    if section_name == "derechos_vulnerados":
        if not _ready("hechos"):
            raise HTTPException(status_code=409, detail="Completa y mejora HECHOS antes de DERECHOS VULNERADOS")

    if section_name == "fundamentos_juridicos":
        if not _ready("hechos") or not _ready("derechos_vulnerados"):
            raise HTTPException(status_code=409, detail="Faltan HECHOS y/o DERECHOS VULNERADOS para FUNDAMENTOS JURÍDICOS")

    if section_name == "fundamentos_de_derecho":
        if not _ready("fundamentos_juridicos"):
            raise HTTPException(status_code=409, detail="Faltan FUNDAMENTOS JURÍDICOS para FUNDAMENTOS DE DERECHO")

    if section_name == "ref":
        if not (_ready("derechos_vulnerados") and _ready("fundamentos_juridicos") and _ready("fundamentos_de_derecho")):
            raise HTTPException(status_code=409, detail="REF requiere D + FJ + FD listos")

# ------------------------------------------------------------
# Cadena automática (ACTIVADA: endpoint específico + pipeline)
# ------------------------------------------------------------

def _build_ctx(conn: sqlite3.Connection, case_id: str) -> Dict[str, Any]:
    cur = conn.cursor()
    hechos = _get_best_text(cur.execute(
        "SELECT * FROM sections WHERE case_id=? AND name='hechos'", (case_id,)).fetchone())
    dv = _get_best_text(cur.execute(
        "SELECT * FROM sections WHERE case_id=? AND name='derechos_vulnerados'", (case_id,)).fetchone())
    fj = _get_best_text(cur.execute(
        "SELECT * FROM sections WHERE case_id=? AND name='fundamentos_juridicos'", (case_id,)).fetchone())
    pruebas = _get_best_text(cur.execute(
        "SELECT * FROM sections WHERE case_id=? AND name='pruebas_y_anexos'", (case_id,)).fetchone())

    acc = _compose_people_inline(conn, case_id, "accionante")
    ads = _compose_people_inline(conn, case_id, "accionado")
    people = "Accionante(s): " + (acc or "-") + "\nAccionado(s): " + (ads or "-")

    derechos_detectados = _detect_rights(" ".join([hechos or "", pruebas or ""]))
    return {
        "hechos": hechos,
        "derechos_vulnerados": dv,
        "fundamentos_juridicos": fj,
        "pruebas": pruebas,
        "personas": people,
        "accionantes_inline": acc,   # <— NUEVO
        "accionados_inline": ads,    # <— NUEVO
        "derechos_detectados_dic": derechos_detectados,
    }

def _improve_store(conn: sqlite3.Connection, case_id: str, name: str, llm=None, retriever=None) -> Dict[str, Any]:
    cur = conn.cursor()
    row = cur.execute("SELECT * FROM sections WHERE case_id=? AND name=?", (case_id, name)).fetchone()
    user_text = (row["user_text"] or "").strip()
    ctx = _build_ctx(conn, case_id)
    ai_text, citations = _llm_improve_for_section(
        name=name, user_text=user_text, ctx=ctx, llm=llm, retriever=retriever
    )
    return _save_section_ai(conn, case_id, name, ai_text, citations)

def _suggest_pretensiones(conn: sqlite3.Connection, case_id: str, llm=None) -> str:
    """Genera sugerencias extra de pretensiones con base en HECHOS (+pret limpias)."""
    if not llm:
        return ""
    cur = conn.cursor()
    hechos = _get_best_text(cur.execute(
        "SELECT * FROM sections WHERE case_id=? AND name='hechos'", (case_id,)).fetchone())
    pret = _get_best_text(cur.execute(
        "SELECT * FROM sections WHERE case_id=? AND name='pretensiones'", (case_id,)).fetchone())
    prompt = PROMPT_SUGIERE_PRETENSIONES.format(hechos=hechos, pret=pret)
    try:
        resp = llm.invoke(prompt)
        txt = getattr(resp, "content", None) or str(resp)
        return (txt or "").strip()
    except Exception:
        return ""

def _chain_autogen(conn: sqlite3.Connection, case_id: str, llm=None, retriever=None) -> Dict[str, str]:
    cur = conn.cursor()
    hechos = _get_best_text(cur.execute(
        "SELECT * FROM sections WHERE case_id=? AND name='hechos'", (case_id,)
    ).fetchone())
    pruebas = _get_best_text(cur.execute(
        "SELECT * FROM sections WHERE case_id=? AND name='pruebas_y_anexos'", (case_id,)
    ).fetchone())

    if not hechos.strip():
        raise HTTPException(status_code=400, detail="Faltan HECHOS (mejorados) para encadenar.")

    # 0) Derechos detectados por diccionario para enriquecer el prompt
    derechos_detectados_dic = _detect_rights(" ".join([hechos or "", pruebas or ""]))

    # 1) Derechos (desde Hechos + lista detectada)
    ctx_d = {"hechos": hechos, "derechos_detectados_dic": derechos_detectados_dic}
    ai_d, _ = _llm_improve_for_section(
        "derechos_vulnerados", user_text="", ctx=ctx_d, llm=llm, retriever=None
    )
    _save_section_ai(conn, case_id, "derechos_vulnerados", ai_d, citations=None)

    # 2) Fundamentos jurídicos (4 subpartes, con contexto completo)
    ctx_fj = {
        "hechos": hechos,
        "derechos_vulnerados": ai_d,
        "pruebas": pruebas,
        "derechos_detectados_dic": derechos_detectados_dic,
    }
    ai_fj, _ = _llm_improve_for_section(
        "fundamentos_juridicos", user_text="", ctx=ctx_fj, llm=llm, retriever=None
    )
    _save_section_ai(conn, case_id, "fundamentos_juridicos", ai_fj, citations=None)

    # 3) Fundamentos de derecho (RAG, desde FJ)
    ctx_fd = {"fundamentos_juridicos": ai_fj, "hechos": hechos}
    ai_fd, cites_fd = _llm_improve_for_section(
        "fundamentos_de_derecho", user_text="", ctx=ctx_fd, llm=llm, retriever=retriever
    )
    _save_section_ai(conn, case_id, "fundamentos_de_derecho", ai_fd, citations=cites_fd)

    # 4) REF (síntesis D + FJ + FD)
    acc_str = _compose_people_inline(conn, case_id, "accionante")
    ads_str = _compose_people_inline(conn, case_id, "accionado")
    ctx_ref = {
        "derechos_vulnerados": ai_d,
        "fundamentos_juridicos": ai_fj,
        "fundamentos_de_derecho": ai_fd,
        "accionantes_inline": acc_str,   # <— NUEVO
        "accionados_inline": ads_str,    # <— NUEVO
    }
    ai_ref, _ = _llm_improve_for_section("ref", user_text="", ctx=ctx_ref, llm=llm, retriever=None)
    _save_section_ai(conn, case_id, "ref", ai_ref, citations=None)


    # Refresca rights_detected simples (diccionario) para panel
    for r in derechos_detectados_dic:
        _set_right(conn, case_id, r)

    return {
        "derechos_vulnerados": ai_d,
        "fundamentos_juridicos": ai_fj,
        "fundamentos_de_derecho": ai_fd,
        "ref": ai_ref,
    }

# ------------------------------------------------------------
# Router factory
# ------------------------------------------------------------

def create_router(
    retriever=None,
    llm=None,
    export_dir: str = EXPORT_DIR_DEFAULT,
    db_path: str = os.path.join(DATA_DIR_DEFAULT, "tutelas.db"),
    top_k_default: Optional[int] = None,
    max_tokens_default: Optional[int] = None,
    **_ignore
) -> APIRouter:
    os.makedirs(export_dir, exist_ok=True)
    _init_db(db_path)
    router = APIRouter()

    # ---------------------- CASES ----------------------------
    @router.post("/case", response_model=CaseCreateResp)
    def create_case():
        conn = _connect(db_path)
        cur = conn.cursor()
        case_id = uuid.uuid4().hex[:12]
        now = _now()
        cur.execute("""
            INSERT INTO cases (id, title, status, created_at, updated_at)
            VALUES (?,?,?,?,?)
        """, (case_id, "Acción de Tutela", "draft", now, now))
        conn.commit()
        _ensure_sections_for_case(conn, case_id)

        # por defecto juramento / ref inicial (se sobreescribirá luego con IA)
        cur.execute("UPDATE sections SET user_text=? WHERE case_id=? AND name='cumplimiento_art_37'",
                    ("JURAMENTO: Manifiesto bajo la gravedad del juramento que no se ha presentado ninguna otra acción de tutela por los mismos hechos y derechos.", case_id))
        cur.execute("UPDATE sections SET user_text=? WHERE case_id=? AND name='ref'",
                    ("Acción de Tutela para proteger el derecho a la salud en conexidad con el derecho a la vida.", case_id))
        conn.commit()
        conn.close()
        return CaseCreateResp(case_id=case_id)

    @router.get("/cases", response_model=List[CaseListItem])
    def list_cases():
        conn = _connect(db_path)
        cur = conn.cursor()
        rows = cur.execute("SELECT id, title, status, updated_at FROM cases ORDER BY updated_at DESC").fetchall()
        conn.close()
        return [CaseListItem(case_id=r["id"], title=r["title"], status=r["status"], updated_at=r["updated_at"]) for r in rows]

    @router.get("/case/{case_id}")
    def get_case(case_id: str):
        conn = _connect(db_path)
        bundle = _get_case_bundle(conn, case_id)
        conn.close()
        return bundle

    # ---------------------- PARTIES --------------------------
    @router.post("/case/{case_id}/party")
    def upsert_party(case_id: str, req: PartyUpsertReq):
        conn = _connect(db_path)
        _get_case_bundle(conn, case_id)
        data = _upsert_party(conn, case_id, req.dict())
        conn.close()
        return data

    # ---------------------- SECTIONS -------------------------
    @router.post("/case/{case_id}/section/{name}")
    def save_section(case_id: str, name: str, req: SectionSaveReq):
        conn = _connect(db_path)
        _get_case_bundle(conn, case_id)  # valida

        # normalización de INTRO (reemplazos por personas)
        if name == "intro" and (req.user_text or "").strip():
            ads = _compose_people_inline(conn, case_id, "accionado")
            if ads:
                txt = req.user_text
                txt = X_RE.sub(ads, txt)
                txt = RE_SENTINEL.sub(ads, txt)
                txt = RE_CONTRA_BLOCK.sub(rf"\1{ads}\3", txt)
                req.user_text = txt

        row = _save_section_user_text(conn, case_id, name, req.user_text or "")

        # --- Mejora automática al guardar + invalidaciones dependientes ---
        if name == "hechos":
            _improve_store(conn, case_id, "hechos", llm=llm, retriever=retriever)
            # Derechos dependen de Hechos; y a su vez FJ, FD, REF dependen en cadena
            _invalidate_sections(conn, case_id, ["derechos_vulnerados", "fundamentos_juridicos", "fundamentos_de_derecho", "ref"])

        elif name == "pretensiones":
            # Mejora + sugerencias (aunque no alimenta Derechos)
            _improve_store(conn, case_id, "pretensiones", llm=llm, retriever=retriever)
            extra = _suggest_pretensiones(conn, case_id, llm=llm)
            if extra.strip():
                cur2 = conn.cursor()
                prev = cur2.execute("SELECT * FROM sections WHERE case_id=? AND name='pretensiones'", (case_id,)).fetchone()
                ai_base = (prev["ai_text"] or prev["user_text"] or "").strip()
                combined = (ai_base + ("\n\nPretensiones sugeridas:\n" + extra)).strip()
                _save_section_ai(conn, case_id, "pretensiones", combined, citations=None)
            # NO invalida cadena (por tu regla, Derechos salen SOLO de Hechos)

        elif name == "pruebas_y_anexos":
            _improve_store(conn, case_id, "pruebas_y_anexos", llm=llm, retriever=retriever)
            # Fundamentos jurídicos (y en consecuencia FD, REF) dependen de Pruebas
            _invalidate_sections(conn, case_id, ["fundamentos_juridicos", "fundamentos_de_derecho", "ref"])

        conn.close()
        return {"row": row, "cascade": []}

    @router.post("/case/{case_id}/section/{name}/improve", response_model=SectionImproveResp)
    def improve_section(case_id: str, name: str):
        conn = _connect(db_path)
        _get_case_bundle(conn, case_id)
        meta = SECTIONS_CONFIG.get(name)
        if not meta:
            conn.close()
            raise HTTPException(status_code=404, detail=f"Sección desconocida: {name}")
        if not meta["needs_llm"]:
            conn.close()
            raise HTTPException(status_code=400, detail="Esta sección no requiere LLM")

        _check_dependencies_or_409(conn, case_id, name)
        updated = _improve_store(conn, case_id, name, llm=llm, retriever=retriever)

        # Si mejoramos derechos → actualizar derechos_detected (encadenes mínimos)
        if name in ("derechos_vulnerados",):
            rights = _detect_rights(updated["ai_text"] or updated["user_text"] or "")
            for r in rights:
                _set_right(conn, case_id, r)

        conn.close()
        return SectionImproveResp(ai_text=updated["ai_text"], citations=json.loads(updated["citations_json"] or "[]"))

    @router.post("/case/{case_id}/section/{name}/approve")
    def approve_section(case_id: str, name: str, req: SectionApproveReq):
        conn = _connect(db_path)
        _get_case_bundle(conn, case_id)
        updated = _approve_section(conn, case_id, name, source=req.source)
        conn.close()
        return updated

    # ---------------------- RIGHTS ---------------------------
    @router.post("/case/{case_id}/rights/detect", response_model=RightsDetectResp)
    def detect_rights(case_id: str):
        conn = _connect(db_path)
        cur = conn.cursor()
        hechos = cur.execute("SELECT * FROM sections WHERE case_id=? AND name='hechos'", (case_id,)).fetchone()
        derechos = cur.execute("SELECT * FROM sections WHERE case_id=? AND name='derechos_vulnerados'", (case_id,)).fetchone()
        text = " ".join([
            (hechos["final_text"] or hechos["ai_text"] or hechos["user_text"] or ""),
            (derechos["final_text"] or derechos["ai_text"] or derechos["user_text"] or ""),
        ])
        rights = _detect_rights(text)
        for r in rights:
            _set_right(conn, case_id, r)
        conn.close()
        return RightsDetectResp(rights=rights)

    @router.post("/case/{case_id}/rights/{right_name}/argue")
    def argue_right(case_id: str, right_name: str):
        conn = _connect(db_path)
        cur = conn.cursor()
        hechos = cur.execute("SELECT * FROM sections WHERE case_id=? AND name='hechos'", (case_id,)).fetchone()
        derechos = cur.execute("SELECT * FROM sections WHERE case_id=? AND name='derechos_vulnerados'", (case_id,)).fetchone()
        user_text = (
            f"Hechos:\n{(hechos['final_text'] or hechos['ai_text'] or hechos['user_text'] or '').strip()}\n\n"
            f"Derechos:\n{(derechos['final_text'] or derechos['ai_text'] or derechos['user_text'] or '').strip()}\n\n"
            f"Derecho específico: {right_name}"
        )
        ai_text, citations = _llm_improve_for_section(
            name="derechos_vulnerados",
            user_text=user_text,
            ctx={},
            llm=llm,
            retriever=retriever
        )
        row = _set_right(conn, case_id, right_name, argument_ai=ai_text, sources=citations)
        conn.close()
        return row

    # Endpoint para refrescar intro manualmente (útil para casos viejos)
    @router.post("/case/{case_id}/intro/refresh")
    def refresh_intro(case_id: str):
        conn = _connect(db_path)
        _get_case_bundle(conn, case_id)
        _refresh_intro_after_party(conn, case_id)
        row = conn.execute("SELECT * FROM sections WHERE case_id=? AND name='intro'", (case_id,)).fetchone()
        conn.close()
        return {"ok": True, "intro": _get_best_text(row)}

    # ---------------------- CADENA (nuevo endpoint) ----------
    @router.post("/case/{case_id}/chain/autogen")
    def chain_autogen(case_id: str):
        conn = _connect(db_path)
        _get_case_bundle(conn, case_id)
        out = _chain_autogen(conn, case_id, llm=llm, retriever=retriever)
        conn.close()
        return {"ok": True, "generated": out}

    # ---------------------- PIPELINE (IA controlada) ---------
    # Ahora el pipeline corre: HECHOS -> (opcional) PRETENSIONES + sugerencias -> CADENA
    @router.post("/case/{case_id}/run-pipeline", response_model=RunPipelineResp)
    def run_pipeline(case_id: str):
        conn = _connect(db_path)
        _get_case_bundle(conn, case_id)
        ran: List[str] = []

        # 1) Hechos
        _improve_store(conn, case_id, "hechos", llm=llm, retriever=retriever)
        ran.append("hechos")

        # 2) Pretensiones (si hay texto del usuario)
        cur = conn.cursor()
        pret_row = cur.execute("SELECT * FROM sections WHERE case_id=? AND name='pretensiones'", (case_id,)).fetchone()
        if (pret_row["user_text"] or "").strip():
            _improve_store(conn, case_id, "pretensiones", llm=llm, retriever=retriever)
            extra = _suggest_pretensiones(conn, case_id, llm=llm)
            if extra.strip():
                prev = cur.execute("SELECT * FROM sections WHERE case_id=? AND name='pretensiones'", (case_id,)).fetchone()
                ai_base = (prev["ai_text"] or prev["user_text"] or "").strip()
                combined = (ai_base + ("\n\nPretensiones sugeridas:\n" + extra)).strip()
                _save_section_ai(conn, case_id, "pretensiones", combined, citations=None)
            ran.append("pretensiones")

        # 3) Cadena (derechos → fundamentos → fundamentos de derecho → ref)
        _chain_autogen(conn, case_id, llm=llm)
        ran.extend(["derechos_vulnerados","fundamentos_juridicos","fundamentos_de_derecho","ref"])

        conn.close()
        return RunPipelineResp(ran=ran)

    # ---------------------- COMPOSE / EXPORT -----------------
    @router.get("/case/{case_id}/compose-final", response_model=ComposeFinalResp)
    def compose_final(case_id: str):
        """Devuelve el documento completo concatenado en texto plano."""
        conn = _connect(db_path)
        _get_case_bundle(conn, case_id)
        text = _compose_full_text(conn, case_id)
        conn.close()
        return ComposeFinalResp(full_text=text)

    # NUEVO: bundle estructurado (útil para render.html editable)
    @router.get("/case/{case_id}/compose-structured")
    def compose_structured(case_id: str):
        conn = _connect(db_path)
        cur = conn.cursor()
        _get_case_bundle(conn, case_id)

        def row(name: str):
            r = cur.execute("SELECT * FROM sections WHERE case_id=? AND name=?", (case_id, name)).fetchone()
            return {
                "user": (r["user_text"] if r else "") or "",
                "ai": (r["ai_text"] if r else "") or "",
                "final": (r["final_text"] if r else "") or "",
                "status": (r["status"] if r else "empty")
            }

        out = {
            "auto": {
                "intro": _get_best_text(cur.execute("SELECT * FROM sections WHERE case_id=? AND name='intro'", (case_id,)).fetchone()),
                "notificaciones": _get_best_text(cur.execute("SELECT * FROM sections WHERE case_id=? AND name='notificaciones'", (case_id,)).fetchone()),
                "firmas": _get_best_text(cur.execute("SELECT * FROM sections WHERE case_id=? AND name='firmas'", (case_id,)).fetchone()),
                "cumplimiento_art_37": _get_best_text(cur.execute("SELECT * FROM sections WHERE case_id=? AND name='cumplimiento_art_37'", (case_id,)).fetchone()),
            },
            "hechos": row("hechos"),
            "pretensiones": row("pretensiones"),
            "pruebas_y_anexos": row("pruebas_y_anexos"),
            "generados": {
                "derechos_vulnerados": row("derechos_vulnerados"),
                "fundamentos_juridicos": row("fundamentos_juridicos"),
                "fundamentos_de_derecho": row("fundamentos_de_derecho"),
                "ref": row("ref"),
            }
        }
        conn.close()
        return out

    # ---------------------- ENSURE (auto por pantalla) -------------------------
    @router.get("/case/{case_id}/ensure/derechos_vulnerados")
    def ensure_derechos(case_id: str):
        conn = _connect(db_path)
        _get_case_bundle(conn, case_id)
        _check_dependencies_or_409(conn, case_id, "derechos_vulnerados")
        updated = _improve_store(conn, case_id, "derechos_vulnerados", llm=llm, retriever=retriever)
        # refresca derechos_detected simples
        rights = _detect_rights((updated["ai_text"] or updated["user_text"] or ""))
        for r in rights:
            _set_right(conn, case_id, r)
        conn.close()
        return {"name": "derechos_vulnerados", "ai_text": updated["ai_text"]}

    @router.get("/case/{case_id}/ensure/fundamentos_juridicos")
    def ensure_fund_j(case_id: str):
        conn = _connect(db_path)
        _get_case_bundle(conn, case_id)
        _check_dependencies_or_409(conn, case_id, "fundamentos_juridicos")
        updated = _improve_store(conn, case_id, "fundamentos_juridicos", llm=llm, retriever=retriever)
        conn.close()
        return {"name": "fundamentos_juridicos", "ai_text": updated["ai_text"]}

    @router.get("/case/{case_id}/ensure/fundamentos_de_derecho")
    def ensure_fund_d(case_id: str):
        conn = _connect(db_path)
        _get_case_bundle(conn, case_id)
        _check_dependencies_or_409(conn, case_id, "fundamentos_de_derecho")
        updated = _improve_store(conn, case_id, "fundamentos_de_derecho", llm=llm, retriever=retriever)
        conn.close()
        return {"name": "fundamentos_de_derecho", "ai_text": updated["ai_text"]}

    @router.get("/case/{case_id}/ensure/ref")
    def ensure_ref(case_id: str):
        conn = _connect(db_path)
        _get_case_bundle(conn, case_id)
        _check_dependencies_or_409(conn, case_id, "ref")
        updated = _improve_store(conn, case_id, "ref", llm=llm, retriever=retriever)
        conn.close()
        return {"name": "ref", "ai_text": updated["ai_text"]}

    @router.post("/case/{case_id}/export-docx", response_model=ExportDocxResp)
    def export_docx(case_id: str):
        conn = _connect(db_path)
        urls = _export_docx(conn, case_id, export_dir)
        conn.close()
        return ExportDocxResp(**urls)

    # NUEVOS shortcuts GET para descargas (útiles en front simple)
    @router.get("/export/docx/{case_id}")
    def export_docx_get(case_id: str):
        conn = _connect(db_path)
        urls = _export_docx(conn, case_id, export_dir)
        conn.close()
        # compat con UIs que esperan 'path'/'filename'
        return {"path": urls["docx_url"], "filename": os.path.basename(urls["docx_url"])}

    @router.get("/export/json/{case_id}")
    def export_json_get(case_id: str):
        # exporta el bundle de caso (no el texto concatenado)
        conn = _connect(db_path)
        urls = _export_docx(conn, case_id, export_dir)
        conn.close()
        return {"path": urls["json_url"], "filename": os.path.basename(urls["json_url"])}

    return router
